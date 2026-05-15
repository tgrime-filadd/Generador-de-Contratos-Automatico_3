import streamlit as st
import pandas as pd
import zipfile
import io
import re
import subprocess
import tempfile
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document

st.set_page_config(page_title="Contratos · Filadd", page_icon="📄", layout="centered")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&family=JetBrains+Mono:wght@400;500&display=swap');
:root {
    --blue:#2424ff; --blue-l:#4444ff; --blue-dim:rgba(36,36,255,0.10); --blue-glow:rgba(36,36,255,0.28); --blue-bdr:rgba(36,36,255,0.35);
    --bg:#F4F4F8; --surface:#FFFFFF; --surf2:#EBEBF0; --border:#DCDCE6;
    --text:#0A0A14; --mid:#4A4A5A; --soft:#8888A0;
    --ok:#16A34A; --ok-bg:rgba(22,163,74,0.08); --ok-bdr:rgba(22,163,74,0.22);
    --warn:#D97706; --warn-bg:rgba(217,119,6,0.08); --warn-bdr:rgba(217,119,6,0.22);
    --err:#DC2626; --err-bg:rgba(220,38,38,0.08); --err-bdr:rgba(220,38,38,0.22);
    --r:12px; --rs:8px;
    --font:'Inter',sans-serif; --mono:'JetBrains Mono',monospace;
    --sh:0 1px 8px rgba(0,0,0,0.07); --sh-b:0 4px 20px rgba(36,36,255,0.18);
}
html,body,[data-testid="stAppViewContainer"],[data-testid="stMain"],[data-testid="block-container"]
    { background:var(--bg) !important; font-family:var(--font) !important; }
[data-testid="stHeader"],[data-testid="stDecoration"],[data-testid="stToolbar"]
    { background:var(--bg) !important; border:none !important; display:block; }
[data-testid="stDecoration"],[data-testid="stToolbar"] { display:none !important; }

/* HEADER */
.fl-hd { display:flex; align-items:center; justify-content:space-between;
    background:var(--blue); border-radius:var(--r); padding:20px 28px; margin-bottom:30px; }
.fl-logo { font-size:1.4rem; font-weight:900; letter-spacing:-0.5px; color:#fff; }
.fl-logo-dim { opacity:.55; }
.fl-badge { background:rgba(255,255,255,0.16); color:#fff; font-size:.68rem; font-weight:700;
    letter-spacing:.10em; text-transform:uppercase; padding:5px 13px; border-radius:100px;
    border:1px solid rgba(255,255,255,.24); }

/* TITLE */
.fl-ttl { font-size:1.85rem; font-weight:900; letter-spacing:-.8px; color:var(--text); margin:0 0 5px; line-height:1.15; }
.fl-ttl em { font-style:normal; color:var(--blue); }
.fl-sub { font-size:.90rem; color:var(--soft); margin-bottom:34px; line-height:1.55; font-weight:400; }

/* STEPS */
.fl-step { display:flex; align-items:center; gap:10px; margin:26px 0 13px; }
.fl-n { width:26px; height:26px; background:var(--blue); color:#fff; font-weight:800; font-size:.76rem;
    border-radius:50%; display:inline-flex; align-items:center; justify-content:center;
    flex-shrink:0; box-shadow:0 2px 8px var(--blue-glow); }
.fl-nt { font-size:.90rem; font-weight:700; color:var(--text); letter-spacing:-.2px; }

/* UPLOADER */
[data-testid="stFileUploader"]
    { background:var(--surface) !important; border:1.5px dashed var(--border) !important;
      border-radius:var(--r) !important; box-shadow:var(--sh) !important;
      transition:border-color .2s,box-shadow .2s !important; }
[data-testid="stFileUploader"]:hover
    { border-color:var(--blue) !important; box-shadow:0 0 0 3px var(--blue-dim) !important; }
[data-testid="stFileUploader"] label
    { color:var(--mid) !important; font-weight:600 !important; font-size:.85rem !important; }
[data-testid="stFileUploaderDropzone"] p,[data-testid="stFileUploaderDropzone"] span
    { color:var(--soft) !important; font-size:.81rem !important; }
[data-testid="stFileUploaderFileName"]
    { color:var(--blue) !important; font-family:var(--mono) !important; font-size:.79rem !important; }

/* CARDS */
.fl-card { background:var(--surface); border:1px solid var(--border); border-radius:var(--r);
    padding:18px 20px; box-shadow:var(--sh); }
.fl-ct { font-size:.68rem; font-weight:700; letter-spacing:.10em; text-transform:uppercase;
    color:var(--soft); margin-bottom:12px; display:flex; align-items:center; gap:7px; }
.fl-ct::before { content:''; display:inline-block; width:6px; height:6px;
    background:var(--blue); border-radius:50%; }
.fl-pv { display:inline-block; background:var(--blue-dim); border:1px solid var(--blue-bdr);
    color:var(--blue) !important; font-family:var(--mono) !important; font-size:.74rem;
    font-weight:500; padding:3px 10px; border-radius:6px; margin:3px 3px 3px 0; }
.fl-pc { display:inline-block; background:var(--surf2); border:1px solid var(--border);
    color:var(--mid) !important; font-family:var(--mono) !important; font-size:.74rem;
    font-weight:500; padding:3px 10px; border-radius:6px; margin:3px 3px 3px 0; }
.fl-rc { font-size:.76rem; color:var(--soft); margin-top:10px; font-weight:500; }

/* BOXES */
.fl-box { border-radius:var(--rs); padding:13px 16px; margin:7px 0;
    font-size:.85rem; line-height:1.55; font-weight:500; font-family:var(--font); }
.fl-ok   { background:var(--ok-bg);   border:1px solid var(--ok-bdr);   color:var(--ok)   !important; }
.fl-wa   { background:var(--warn-bg); border:1px solid var(--warn-bdr); color:var(--warn) !important; }
.fl-er   { background:var(--err-bg);  border:1px solid var(--err-bdr);  color:var(--err)  !important; }
.fl-in   { background:var(--blue-dim);border:1px solid var(--blue-bdr); color:var(--blue) !important; }
.fl-box b   { font-weight:700; }
.fl-box code { background:rgba(0,0,0,.07) !important; color:inherit !important;
    font-family:var(--mono) !important; font-size:.79rem; padding:1px 6px; border-radius:4px; }
.fl-box a { color:inherit !important; }

/* RESULT GRID */
.fl-rg { display:grid; grid-template-columns:repeat(3,1fr); gap:10px; margin:16px 0 20px; }
.fl-rc2 { background:var(--surface); border:1px solid var(--border); border-radius:var(--r);
    padding:16px; text-align:center; box-shadow:var(--sh); }
.fl-rc2.hl { background:var(--blue-dim); border-color:var(--blue-bdr); }
.fl-rn { font-size:2rem; font-weight:900; letter-spacing:-1px; color:var(--text); line-height:1; display:block; }
.fl-rc2.hl .fl-rn { color:var(--blue) !important; }
.fl-rl { font-size:.67rem; color:var(--soft); text-transform:uppercase;
    letter-spacing:.07em; font-weight:700; margin-top:4px; display:block; }

/* DIVIDER */
.fl-div { height:1px; background:var(--border); margin:26px 0; border:none; }

/* BUTTONS */
.stButton>button { font-family:var(--font) !important; font-weight:700 !important;
    font-size:.87rem !important; border-radius:var(--rs) !important; width:100% !important;
    transition:all .15s ease !important; letter-spacing:-.1px !important; padding:.6rem 1.4rem !important; }
button[data-testid="baseButton-primary"],.stButton>button[kind="primary"]
    { background:var(--blue) !important; color:#fff !important; border:none !important;
      box-shadow:var(--sh-b) !important; }
button[data-testid="baseButton-primary"]:hover,.stButton>button[kind="primary"]:hover
    { background:var(--blue-l) !important; transform:translateY(-1px) !important;
      box-shadow:0 6px 26px var(--blue-glow) !important; }
button[data-testid="baseButton-secondary"],.stButton>button[kind="secondary"]
    { background:var(--surface) !important; color:var(--text) !important;
      border:1.5px solid var(--border) !important; box-shadow:var(--sh) !important; }
button[data-testid="baseButton-secondary"]:hover,.stButton>button[kind="secondary"]:hover
    { border-color:var(--blue) !important; color:var(--blue) !important;
      box-shadow:0 0 0 3px var(--blue-dim) !important; transform:translateY(-1px) !important; }
.stButton>button:disabled
    { background:var(--surf2) !important; color:var(--soft) !important;
      border:1px solid var(--border) !important; box-shadow:none !important;
      transform:none !important; opacity:.55 !important; }
[data-testid="stDownloadButton"]>button
    { background:var(--blue) !important; color:#fff !important; font-weight:800 !important;
      border:none !important; font-size:.92rem !important; box-shadow:var(--sh-b) !important;
      width:100% !important; padding:.76rem 1.4rem !important; border-radius:var(--rs) !important; }
[data-testid="stDownloadButton"]>button:hover
    { background:var(--blue-l) !important; transform:translateY(-2px) !important;
      box-shadow:0 10px 32px var(--blue-glow) !important; }

/* PROGRESS */
[data-testid="stProgressBar"]>div
    { background:var(--surf2) !important; border-radius:6px !important; overflow:hidden !important; }
[data-testid="stProgressBar"]>div>div
    { background:var(--blue) !important; border-radius:6px !important; box-shadow:0 0 8px var(--blue-glow) !important; }

/* EXPANDER */
[data-testid="stExpander"]
    { background:var(--surface) !important; border:1px solid var(--border) !important;
      border-radius:var(--r) !important; overflow:hidden !important; box-shadow:var(--sh) !important; }
[data-testid="stExpander"] summary
    { background:var(--surface) !important; color:var(--text) !important;
      font-weight:600 !important; font-family:var(--font) !important;
      font-size:.87rem !important; padding:14px 18px !important; }
[data-testid="stExpander"] summary:hover { color:var(--blue) !important; }

/* DATAFRAME */
[data-testid="stDataFrame"]
    { border:1px solid var(--border) !important; border-radius:var(--r) !important;
      overflow:hidden !important; box-shadow:var(--sh) !important; }

/* CODE */
code { background:var(--surf2) !important; color:var(--blue) !important;
    font-family:var(--mono) !important; font-size:.79rem !important;
    padding:2px 6px !important; border-radius:4px !important; border:1px solid var(--border) !important; }

/* MARKDOWN */
[data-testid="stMarkdownContainer"] p
    { color:var(--soft) !important; font-size:.87rem !important; line-height:1.55 !important; }
[data-testid="stMarkdownContainer"] strong { color:var(--text) !important; }
[data-testid="stMarkdownContainer"] li     { color:var(--mid) !important; font-size:.85rem !important; }
[data-testid="stMarkdownContainer"] table
    { border-collapse:collapse !important; width:100% !important; font-size:.82rem !important;
      border-radius:var(--rs) !important; overflow:hidden !important; border:1px solid var(--border) !important; }
[data-testid="stMarkdownContainer"] th
    { background:var(--blue-dim) !important; color:var(--blue) !important; font-weight:700 !important;
      padding:8px 12px !important; border:1px solid var(--blue-bdr) !important;
      font-family:var(--mono) !important; font-size:.77rem !important; }
[data-testid="stMarkdownContainer"] td
    { padding:7px 12px !important; border:1px solid var(--border) !important; color:var(--mid) !important; }
[data-testid="stMarkdownContainer"] tr:nth-child(even) td { background:var(--surf2) !important; }

/* SCROLLBAR */
::-webkit-scrollbar { width:5px; height:5px; }
::-webkit-scrollbar-track { background:var(--bg); }
::-webkit-scrollbar-thumb { background:var(--border); border-radius:3px; }
::-webkit-scrollbar-thumb:hover { background:var(--blue); }

/* FOOTER */
.fl-ft { text-align:center; font-size:.74rem; color:var(--soft);
    padding:32px 0 16px; font-weight:500; font-family:var(--font); }
.fl-ft a { color:var(--blue) !important; text-decoration:none; }
</style>
""", unsafe_allow_html=True)


# ─── FUNCIONES ────────────────────────────────────────────────────────────────

def extract_variables_from_docx(docx_file) -> set:
    docx_file.seek(0)
    doc = Document(docx_file)
    variables, pattern = set(), re.compile(r'\{\{(\w+)\}\}')
    for para in doc.paragraphs:
        variables.update(pattern.findall(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    variables.update(pattern.findall(para.text))
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            for para in hdr.paragraphs:
                variables.update(pattern.findall(para.text))
    return variables


def validate_data(template_vars: set, df: pd.DataFrame) -> dict:
    errors, warnings = [], []
    excel_cols = set(df.columns.str.strip().str.upper())
    tvars = {v.upper() for v in template_vars}
    missing = tvars - excel_cols
    if missing:
        errors.append(f"Variables <code>{'</code>, <code>'.join(sorted(missing))}</code> no existen en el Excel.")
    df2 = df.copy(); df2.columns = df2.columns.str.strip().str.upper()
    for col in list(tvars & excel_cols):
        empty = df2[df2[col].isna() | (df2[col].astype(str).str.strip() == '')].index.tolist()
        if empty:
            warnings.append(f"Columna <code>{col}</code>: celdas vacías en filas {[r+2 for r in empty]}")
    id_kw   = ['DNI','RUT','CPF','CNPJ','CC','CEDULA','DOCUMENTO','IDENTIFICACION','NIT','PASSPORT','PASAPORTE']
    excl_kw = ['MONTO','VALOR','SALARIO','SUELDO','PRECIO','AMOUNT','SALARY','WAGE',
                'NACIONAL','COUNTRY','PAIS','CIDADE','CIUDAD','CARGO','OCUPACION',
                'GENERO','SEXO','ESTADO','STATUS','TIPO','TYPE','NIVEL','QUANTIDADE','QUANTITY']
    for col in df2.columns:
        col_up = col.upper()
        if any(e in col_up for e in excl_kw):
            continue
        if any(k in col_up for k in id_kw):
            dup = df2[df2[col].duplicated(keep=False) & df2[col].notna()]
            if not dup.empty:
                warnings.append(f"Columna <code>{col}</code>: duplicados → {dup[col].unique().tolist()}")
    return {"errors": errors, "warnings": warnings, "valid": len(errors) == 0}


def filter_df(df: pd.DataFrame, template_vars: set):
    df_n = df.copy(); df_n.columns = df_n.columns.str.strip().str.upper()
    tvars = {v.upper() for v in template_vars}
    common = list(tvars & set(df_n.columns))
    nk = next((c for c in df_n.columns if 'NOMBRE' in c or 'NAME' in c), df_n.columns[0])
    skip_m, skip_d, keep = [], [], []
    for idx, row in df_n.iterrows():
        empty = [c for c in common if pd.isna(row.get(c)) or str(row.get(c,'')).strip()=='']
        if empty: skip_m.append(f"Fila {idx+2} — {row.get(nk,'?')}: falta {', '.join(empty)}")
        else: keep.append(idx)
    df_c = df_n.loc[keep].copy()
    id_kw   = ['DNI','RUT','CPF','CNPJ','CC','CEDULA','DOCUMENTO','IDENTIFICACION','NIT','PASSPORT','PASAPORTE']
    excl_kw = ['MONTO','VALOR','SALARIO','SUELDO','PRECIO','AMOUNT','SALARY','WAGE',
                'NACIONAL','COUNTRY','PAIS','CIDADE','CIUDAD','CARGO','OCUPACION',
                'GENERO','SEXO','ESTADO','STATUS','TIPO','TYPE','NIVEL','QUANTIDADE','QUANTITY']
    dup_rows_idx = set()
    for ic in df_c.columns:
        ic_up = ic.upper()
        if any(e in ic_up for e in excl_kw):
            continue
        if any(k in ic_up for k in id_kw):
            mask = df_c[ic].duplicated(keep='first') & df_c[ic].notna()
            for idx, row in df_c[mask].iterrows():
                skip_d.append(f"Fila {idx+2} — {row.get(nk,'?')}: {ic} duplicado ({row[ic]})")
                dup_rows_idx.add(idx)
    df_dup = df_c.loc[list(dup_rows_idx)].copy() if dup_rows_idx else pd.DataFrame()
    df_c = df_c[~df_c.index.isin(dup_rows_idx)]
    return df_c, skip_m, skip_d, df_dup


def to_pdf(docx_bytes: io.BytesIO):
    with tempfile.TemporaryDirectory() as td:
        dp = Path(td)/"c.docx"; pp = Path(td)/"c.pdf"
        with open(dp,'wb') as f: docx_bytes.seek(0); f.write(docx_bytes.read())
        for cmd in ['libreoffice','soffice']:
            try:
                subprocess.run([cmd,'--headless','--convert-to','pdf','--outdir',td,str(dp)],
                               capture_output=True,text=True,timeout=60)
                if pp.exists(): return open(pp,'rb').read()
            except (FileNotFoundError,subprocess.TimeoutExpired): continue
        try:
            from docx2pdf import convert; convert(str(dp),str(pp))
            if pp.exists(): return open(pp,'rb').read()
        except Exception: pass
    return None


def check_pdf():
    for cmd in ['libreoffice','soffice']:
        try: subprocess.run([cmd,'--version'],capture_output=True,timeout=5); return True, cmd
        except (FileNotFoundError,subprocess.TimeoutExpired): continue
    try: import docx2pdf; return True,"docx2pdf"
    except ImportError: pass
    return False,""


# ─── SESSION STATE ─────────────────────────────────────────────────────────────
for k,v in [('validated',False),('validation_result',None),('template_vars',set()),('df',None)]:
    if k not in st.session_state: st.session_state[k] = v

pdf_ok, pdf_tool = check_pdf()

# ─── HEADER ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="fl-hd">
  <div class="fl-logo">filadd<span class="fl-logo-dim">.</span></div>
  <div class="fl-badge">Herramienta interna</div>
</div>
<p class="fl-ttl">Generador de <em>Contratos</em></p>
<p class="fl-sub">Carga una plantilla Word y un Excel con datos — el sistema genera todos los contratos en segundos.</p>
""", unsafe_allow_html=True)

# ─── PASO 1 ─────────────────────────────────────────────────────────────────
st.markdown('<div class="fl-step"><span class="fl-n">1</span><span class="fl-nt">Sube los archivos</span></div>', unsafe_allow_html=True)
c1, c2 = st.columns(2)
with c1: template_file = st.file_uploader("Plantilla Word (.docx)", type=["docx"], help="Variables: {{NOMBRE}}, {{DNI}}, {{MONTO}}…")
with c2: excel_file    = st.file_uploader("Datos Excel (.xlsx)", type=["xlsx"], help="Encabezados = nombres de las variables")

# ─── FLUJO PRINCIPAL ────────────────────────────────────────────────────────
if template_file and excel_file:

    # PASO 2 – Vista previa
    st.markdown('<div class="fl-step"><span class="fl-n">2</span><span class="fl-nt">Vista previa</span></div>', unsafe_allow_html=True)
    ca, cb = st.columns(2)

    with ca:
        try:
            vars_found = extract_variables_from_docx(template_file)
            st.session_state.template_vars = vars_found
            pills = "".join(f'<span class="fl-pv">{{{{{v}}}}}</span>' for v in sorted(vars_found))
            body  = pills if vars_found else '<div class="fl-box fl-wa">⚠️ No hay variables <code>{{...}}</code>.</div>'
            st.markdown(f'<div class="fl-card"><div class="fl-ct">Variables en plantilla</div>{body}</div>', unsafe_allow_html=True)
        except Exception as e:
            st.markdown(f'<div class="fl-box fl-er">❌ Error leyendo plantilla: {e}</div>', unsafe_allow_html=True)

    with cb:
        try:
            df = pd.read_excel(excel_file, dtype=str)
            st.session_state.df = df
            pills = "".join(f'<span class="fl-pc">{c.strip().upper()}</span>' for c in df.columns)
            st.markdown(f'<div class="fl-card"><div class="fl-ct">Columnas en Excel</div>{pills}<div class="fl-rc">{len(df)} fila(s)</div></div>', unsafe_allow_html=True)
        except Exception as e:
            st.markdown(f'<div class="fl-box fl-er">❌ Error leyendo Excel: {e}</div>', unsafe_allow_html=True)

    with st.expander("Ver tabla de datos completa"):
        st.dataframe(df, use_container_width=True)

    st.markdown('<div class="fl-div"></div>', unsafe_allow_html=True)

    # PASO 3 – Validar
    st.markdown('<div class="fl-step"><span class="fl-n">3</span><span class="fl-nt">Validar datos</span></div>', unsafe_allow_html=True)

    if st.button("Validar datos", type="secondary"):
        with st.spinner("Analizando…"):
            res = validate_data(st.session_state.template_vars, st.session_state.df)
            st.session_state.validation_result = res
            st.session_state.validated = res["valid"]

    if st.session_state.validation_result:
        r = st.session_state.validation_result
        for e in r["errors"]:   st.markdown(f'<div class="fl-box fl-er">❌ {e}</div>', unsafe_allow_html=True)
        for w in r["warnings"]: st.markdown(f'<div class="fl-box fl-wa">⚠️ {w}</div>', unsafe_allow_html=True)
        if r["valid"] and not r["warnings"]:
            st.markdown('<div class="fl-box fl-ok">✅ Todo correcto — puedes generar los contratos.</div>', unsafe_allow_html=True)
        elif r["valid"]:
            st.markdown('<div class="fl-box fl-wa">⚠️ Aprobado con advertencias. Las filas con problemas serán omitidas automáticamente.</div>', unsafe_allow_html=True)

    st.markdown('<div class="fl-div"></div>', unsafe_allow_html=True)

    # PASO 4 – Generar
    st.markdown('<div class="fl-step"><span class="fl-n">4</span><span class="fl-nt">Generar contratos</span></div>', unsafe_allow_html=True)

    if not pdf_ok:
        st.markdown(f'<div class="fl-box fl-wa">⚠️ <b>LibreOffice no detectado.</b> Los contratos se generarán en <code>.docx</code>. Instala <a href="https://www.libreoffice.org/download/libreoffice/" target="_blank">LibreOffice</a> para obtener PDFs.</div>', unsafe_allow_html=True)
        output_format = "docx"
    else:
        st.markdown(f'<div class="fl-box fl-in">📄 Soporte PDF detectado <code>({pdf_tool})</code> — los contratos se exportarán como PDF.</div>', unsafe_allow_html=True)
        output_format = "pdf"

    can = bool(st.session_state.validated or (st.session_state.validation_result and st.session_state.validation_result["valid"]))
    n_rows = len(st.session_state.df) if st.session_state.df is not None else 0

    if not can and st.session_state.validation_result is None:
        st.markdown('<div class="fl-box fl-in">ℹ️ Presiona <b>Validar datos</b> antes de continuar.</div>', unsafe_allow_html=True)
    elif not can:
        st.markdown('<div class="fl-box fl-er">❌ Corrige los errores antes de generar.</div>', unsafe_allow_html=True)

    if st.button(f"Generar contratos  ·  {n_rows} registros", type="primary", disabled=not can):
        prog = st.progress(0, text="Preparando…")
        try:
            df_c, sm, sd, df_dup = filter_df(st.session_state.df, st.session_state.template_vars)
            n_total, n_gen = len(st.session_state.df), len(df_c)

            if n_gen == 0:
                st.markdown('<div class="fl-box fl-er">❌ No hay filas válidas. Revisa los datos.</div>', unsafe_allow_html=True)
                prog.empty()
            else:
                if sm:
                    with st.expander(f"❌ {len(sm)} fila(s) omitidas — datos faltantes", expanded=True):
                        for m in sm: st.markdown(f'<div class="fl-box fl-er">⛔ {m}</div>', unsafe_allow_html=True)
                if sd:
                    with st.expander(f"⚠️ {len(sd)} fila(s) omitidas — ID duplicado", expanded=True):
                        for m in sd: st.markdown(f'<div class="fl-box fl-wa">⚠️ {m}</div>', unsafe_allow_html=True)
                        if not df_dup.empty:
                            # Generate zip with duplicate contracts
                            zb_dup = io.BytesIO()
                            with zipfile.ZipFile(zb_dup, 'w', zipfile.ZIP_DEFLATED) as zf_dup:
                                nk_dup = next((c for c in df_dup.columns if 'NOMBRE' in c or 'NAME' in c or 'NOME' in c), df_dup.columns[0])
                                for i, (_, row) in enumerate(df_dup.iterrows()):
                                    nom = row.get(nk_dup, f"#{i+1}")
                                    ctx = {col:(str(val) if pd.notna(val) else '') for col,val in row.items()}
                                    template_file.seek(0)
                                    tpl = DocxTemplate(template_file); tpl.render(ctx)
                                    db = io.BytesIO(); tpl.save(db); db.seek(0)
                                    fn = f"duplicado_{str(nom).replace(' ','_').replace('//','-')}"
                                    if output_format == "pdf":
                                        pb = to_pdf(db)
                                        if pb: zf_dup.writestr(f"{fn}.pdf", pb)
                                        else:  db.seek(0); zf_dup.writestr(f"{fn}.docx", db.read())
                                    else:
                                        zf_dup.writestr(f"{fn}.docx", db.read())
                            zb_dup.seek(0)
                            st.download_button(
                                label=f"⬇  Descargar duplicados.zip  ({len(df_dup)} archivos)",
                                data=zb_dup.getvalue(),
                                file_name="contratos_duplicados.zip",
                                mime="application/zip",
                                help="Contratos generados igualmente para las filas marcadas como duplicadas"
                            )

                nk = next((c for c in df_c.columns if 'NOMBRE' in c or 'NAME' in c), df_c.columns[0])
                zb = io.BytesIO()
                with zipfile.ZipFile(zb, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for i, (_, row) in enumerate(df_c.iterrows()):
                        nom = row.get(nk, f"#{i+1}")
                        prog.progress((i+1)/n_gen, text=f"Generando {i+1}/{n_gen} — {nom}")
                        ctx = {col:(str(val) if pd.notna(val) else '') for col,val in row.items()}
                        template_file.seek(0)
                        tpl = DocxTemplate(template_file); tpl.render(ctx)
                        db = io.BytesIO(); tpl.save(db); db.seek(0)
                        fn = f"contrato_{str(nom).replace(' ','_').replace('//','-')}"
                        if output_format == "pdf":
                            pb = to_pdf(db)
                            if pb: zf.writestr(f"{fn}.pdf", pb)
                            else:  db.seek(0); zf.writestr(f"{fn}.docx", db.read())
                        else:
                            zf.writestr(f"{fn}.docx", db.read())

                prog.progress(1.0, text="¡Listo!")
                zb.seek(0)
                ext = "PDF" if output_format=="pdf" else "DOCX"

                st.markdown(f"""
                <div class="fl-rg">
                  <div class="fl-rc2 hl"><span class="fl-rn">{n_gen}</span><span class="fl-rl">Generados ({ext})</span></div>
                  <div class="fl-rc2"><span class="fl-rn">{len(sd)}</span><span class="fl-rl">Duplicados (omitidos)</span></div>
                  <div class="fl-rc2"><span class="fl-rn">{len(sm)}</span><span class="fl-rl">Datos faltantes</span></div>
                </div>""", unsafe_allow_html=True)

                st.download_button(
                    label=f"⬇  Descargar contratos.zip  ({n_gen} archivos)",
                    data=zb.getvalue(), file_name="contratos_generados.zip", mime="application/zip"
                )
        except Exception as e:
            st.markdown(f'<div class="fl-box fl-er">❌ Error inesperado: {e}</div>', unsafe_allow_html=True)
            st.exception(e)

# ─── ESTADO VACÍO ──────────────────────────────────────────────────────────────
else:
    st.markdown('<div class="fl-box fl-in">ℹ️ <b>Para comenzar:</b> sube tu plantilla <code>.docx</code> y tu archivo <code>.xlsx</code> arriba.</div>', unsafe_allow_html=True)

    with st.expander("¿Cómo preparar la plantilla Word?"):
        st.markdown("""
Escribe las variables con doble llave en cualquier parte del documento:

```
El trabajador {{NOMBRE}}, identificado con {{DNI}},
domiciliado en {{DIRECCION}}, ocupará el cargo de
{{CARGO}} con remuneración {{MONTO}} desde el {{FECHA}}.
```

| Variable | Descripción |
|----------|-------------|
| `{{NOMBRE}}` | Nombre completo |
| `{{DNI}}` | DNI / RUT / CPF / CC |
| `{{DIRECCION}}` | Dirección |
| `{{CARGO}}` | Cargo / Puesto |
| `{{MONTO}}` | Remuneración |
| `{{FECHA}}` | Fecha de inicio |

---

**Las variables pueden tener cualquier nombre**, no solo palabras sueltas. Para nombres compuestos, usa **guión bajo** `_` en lugar de espacios:

| ✅ Correcto | ❌ No funciona |
|------------|---------------|
| `{{FECHA_DE_INICIO}}` | `{{FECHA DE INICIO}}` |
| `{{NOMBRE_COMPLETO}}` | `{{NOMBRE COMPLETO}}` |
| `{{EQUIPO_DE_FUTBOL_FAVORITO}}` | `{{EQUIPO DE FUTBOL FAVORITO}}` |
| `{{NUMERO_DE_CUENTA}}` | `{{NUMERO DE CUENTA}}` |

Recuerda que la columna en el Excel debe tener **exactamente el mismo nombre** que la variable (sin llaves). Por ejemplo, si la plantilla tiene `{{FECHA_DE_INICIO}}`, la columna en el Excel debe llamarse `FECHA_DE_INICIO`.
""")

    with st.expander("¿Cómo preparar el Excel?"):
        st.markdown("""
La primera fila debe contener los mismos nombres que las variables (sin llaves):

| NOMBRE | DNI | DIRECCION | CARGO | MONTO | FECHA |
|--------|-----|-----------|-------|-------|-------|
| Juan Pérez | 12.345.678-9 | Av. Providencia 1234 | Analista | $1.500.000 | 01/04/2025 |
| María López | 98.765.432-1 | Calle 80 #45, Bogotá | Coordinadora | COP 4.200.000 | 01/04/2025 |

Cada fila genera un contrato independiente.
""")

# ─── FOOTER ────────────────────────────────────────────────────────────────────
st.markdown('<div class="fl-ft"><a href="https://filadd.com">filadd.com</a> &nbsp;·&nbsp; Herramienta interna de operaciones</div>', unsafe_allow_html=True)
